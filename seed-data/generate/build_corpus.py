"""
Builds the full synthetic Personal Auto insurance corpus for the Domain Copilot assessment
(variant D2 — insurance claims adjudication, T6 — document in/out).

Usage (from seed-data/generate/, with the local venv active):
    python build_corpus.py

Produces seed-data/corpus/{policy-forms,declarations,endorsements,claims,reference}/... plus
seed-data/corpus/manifest.json describing every document. Re-running regenerates everything from
scratch (idempotent by construction — it's a build step, not a diff).
"""

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw

import facts
from md_to_docx import write_docx
from scanned_pdf import build_scanned_pdf, paginate

ROOT = Path(__file__).parent
CONTENT = ROOT / "content"
OUT = ROOT / "out"
CORPUS = ROOT.parent / "corpus"

manifest: list[dict] = []


def reset_dirs() -> None:
    for d in [OUT, CORPUS]:
        if d.exists():
            for f in d.rglob("*"):
                if f.is_file():
                    f.unlink()
        d.mkdir(parents=True, exist_ok=True)
    for sub in ["policy-forms", "declarations", "endorsements", "claims", "reference"]:
        (CORPUS / sub).mkdir(parents=True, exist_ok=True)


def convert_docx_to_pdf(docx_paths: list[Path], out_dir: Path) -> None:
    if not docx_paths:
        return
    subprocess.run(
        [
            "soffice", "--headless", "--norestore",
            "--convert-to", "pdf", "--outdir", str(out_dir),
            *[str(p) for p in docx_paths],
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
    )


def add_doc(doc_id: str, filename: str, category: str, fmt: str, **extra) -> None:
    manifest.append({"id": doc_id, "filename": filename, "category": category, "format": fmt, **extra})


# --- 1. Hand-authored prose documents (policy forms, guidelines, reference) --------------------

PROSE_DOCS = [
    ("policy_wording_v1", "Meridian Mutual Personal Auto Policy Wording — PAP-2024-STD", "policy-forms", facts.POLICY_FORM_V1),
    ("policy_wording_v2", "Meridian Mutual Personal Auto Policy Wording — PAP-2025-STD", "policy-forms", facts.POLICY_FORM_V2),
    ("exclusions_v1", "Standard Exclusions Addendum — PAP-EXCL-2024", "policy-forms", facts.POLICY_FORM_V1),
    ("exclusions_v2", "Standard Exclusions Addendum — PAP-EXCL-2025", "policy-forms", facts.POLICY_FORM_V2),
    ("adjudication_guidelines", "Claims Adjudication Guidelines", "reference", None),
    ("regulatory_reference", "Ohio Regulatory Reference", "reference", None),
    ("glossary", "Glossary of Insurance Terms", "reference", None),
    ("claims_faq", "Internal Memo: Frequently Asked Adjudication Questions", "reference", None),
    ("underwriting_guidelines", "Personal Auto Underwriting Guidelines", "reference", None),
    ("siu_fraud_indicators", "Special Investigations Unit — Fraud Indicators Guide", "reference", None),
    ("total_loss_methodology", "Total Loss Valuation Methodology", "reference", None),
    ("complaint_handling", "Complaint Handling Procedure", "reference", None),
    ("producer_manual_excerpt", "Producer Manual Excerpt — Binding Authority", "reference", None),
    ("ohio_total_loss_title", "Ohio Total Loss and Salvage Title Requirements", "reference", None),
    ("training_scenarios", "Claims Adjudication Training Scenarios", "reference", None),
    ("version_comparison", "Policy Form Version Comparison — PAP-2024-STD vs. PAP-2025-STD", "reference", None),
    ("pii_data_handling", "Data Privacy and PII Handling Policy", "reference", None),
    ("subrogation_guide", "Subrogation and Litigation Hold Guide", "reference", None),
    ("cat_claims_protocol", "Catastrophe (CAT) Claims Surge Protocol", "reference", None),
    ("unfair_claims_practices", "Ohio Unfair Claims Settlement Practices Reference", "reference", None),
    ("telematics_program", "Telematics and Usage-Based Insurance Program Terms", "reference", None),
    ("claims_handling_manual", "Personal Auto Claims Handling Procedures Manual", "reference", None),
    ("rating_and_rules_manual", "Personal Auto Rating and Rules Manual", "reference", None),
    ("regulatory_compendium", "Ohio Auto Insurance Regulatory Compendium", "reference", None),
    ("claim_scenario_workbook", "Claim Scenario Workbook — Extended Worked Examples", "reference", None),
    ("coverage_interpretation_notes", "Coverage Interpretation Notes — Frequently Disputed Provisions", "reference", None),
    ("reserving_guide", "Claims Financial Reserving Guide", "reference", None),
    ("vendor_network_guide", "Vendor and Repair Network Management Guide", "reference", None),
    ("quality_audit_program", "Claims Quality Assurance and Audit Program", "reference", None),
    ("agent_appointment_manual", "Agent Appointment and Errors & Omissions Reference", "reference", None),
    ("rental_reimbursement_guide", "Rental Reimbursement and Loss of Use Claims Guide", "reference", None),
    ("litigation_referral_guide", "Claims Litigation and Legal Referral Guide", "reference", None),
    ("diminished_value_guide", "Diminished Value Claims Reference", "reference", None),
    ("um_uim_claims_guide", "Uninsured and Underinsured Motorist Claims Handling Guide", "reference", None),
    ("salvage_disposition_guide", "Salvage and Total Loss Vehicle Disposition Guide", "reference", None),
    ("claim_reopening_appeals_guide", "Claim Reopening and Internal Appeals Guide", "reference", None),
    ("roadside_towing_guide", "Roadside Assistance and Towing Coverage Guide", "reference", None),
    ("customer_communication_standards", "Claims Customer Communication Standards", "reference", None),
    ("non_owned_vehicle_guide", "Non-Owned and Borrowed Vehicle Coverage Guide", "reference", None),
    ("glass_windshield_claims_guide", "Glass and Windshield Claims Guide", "reference", None),
    ("multi_vehicle_policy_guide", "Multi-Vehicle Policy and Household Structure Guide", "reference", None),
    ("weather_catastrophe_claims_guide", "Weather-Related and Non-Declared Catastrophe Claims Guide", "reference", None),
    ("new_vehicle_replacement_guide", "New Vehicle Replacement Cost Coverage Guide", "reference", None),
    ("med_pay_pip_coordination_guide", "Medical Payments Coverage and Coordination Guide", "reference", None),
    ("policy_cancellation_nonrenewal_guide", "Policy Cancellation and Non-Renewal Reference", "reference", None),
    ("claims_data_privacy_retention_guide", "Claims Data Retention and Access Control Guide", "reference", None),
    ("claim_intake_triage_guide", "Claim Intake and Triage Guide", "reference", None),
    ("deductible_selection_reference", "Deductible Selection and Application Reference", "reference", None),
    ("rideshare_delivery_use_guide", "Rideshare and Delivery Use Coverage Guide", "reference", None),
]


def build_prose_docs() -> None:
    docx_paths = []
    for slug, title, category, form_version in PROSE_DOCS:
        md_text = (CONTENT / f"{slug}.md").read_text(encoding="utf-8")
        docx_path = OUT / f"{slug}.docx"
        write_docx(md_text, title, docx_path)
        docx_paths.append(docx_path)

    convert_docx_to_pdf(docx_paths, OUT)

    for slug, title, category, form_version in PROSE_DOCS:
        pdf_src = OUT / f"{slug}.pdf"
        pdf_dst = CORPUS / category / f"{slug}.pdf"
        pdf_dst.write_bytes(pdf_src.read_bytes())
        extra = {"formVersion": form_version} if form_version else {}
        add_doc(slug, str(pdf_dst.relative_to(CORPUS)), category, "pdf", title=title, **extra)


# --- 2. Declarations pages (one per policyholder, DOCX) -----------------------------------------

def build_declarations() -> None:
    for ph in facts.POLICYHOLDERS:
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading(f"{facts.INSURER_NAME} — Declarations Page", level=0)
        doc.add_paragraph(f"Policy Number: {ph.policy_number}")
        doc.add_paragraph(f"Policy Form: {ph.form_version} (effective {facts.POLICY_FORM_V1_EFFECTIVE if ph.form_version == facts.POLICY_FORM_V1 else facts.POLICY_FORM_V2_EFFECTIVE})")
        doc.add_paragraph(f"Named Insured: {ph.named_insured}")
        doc.add_paragraph(f"Address: {ph.address}")
        doc.add_paragraph(f"Policy Period Begins: {ph.effective_date}")

        doc.add_heading("Covered Vehicle", level=2)
        doc.add_paragraph(f"{ph.vehicle_year} {ph.vehicle_make} {ph.vehicle_model} — VIN {ph.vin}")

        doc.add_heading("Coverage Selections", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Coverage", "Limit", "Deductible"

        def row(coverage: str, limit: str, deductible: str) -> None:
            cells = table.add_row().cells
            cells[0].text, cells[1].text, cells[2].text = coverage, limit, deductible

        row("Liability (BI per person / BI per accident / PD)", ph.liability_summary + " (thousands)", "N/A")
        if ph.med_pay:
            row("Medical Payments", f"${ph.med_pay:,}", "N/A")
        row("Uninsured/Underinsured Motorist", f"${ph.um_uim_per_person:,} / ${ph.um_uim_per_accident:,}", "N/A")
        if ph.has_collision:
            row("Collision", "Actual Cash Value", f"${ph.collision_deductible:,}")
        else:
            row("Collision", "Not selected", "—")
        if ph.has_comprehensive:
            row("Comprehensive", "Actual Cash Value", f"${ph.comprehensive_deductible:,}")
        if ph.rental_reimbursement_daily:
            row("Transportation Expenses", f"${ph.rental_reimbursement_daily}/day, max 30 days", "N/A")

        if ph.endorsements:
            doc.add_heading("Endorsements Attached", level=2)
            for e in ph.endorsements:
                doc.add_paragraph(e, style="List Bullet")

        slug = f"declarations_{ph.policy_number.replace('-', '_').lower()}"
        docx_path = OUT / f"{slug}.docx"
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_path))
        dst = CORPUS / "declarations" / f"{slug}.docx"
        dst.write_bytes(docx_path.read_bytes())
        add_doc(
            slug, str(dst.relative_to(CORPUS)), "declarations", "docx",
            title=f"Declarations Page — {ph.policy_number}",
            policyNumber=ph.policy_number, formVersion=ph.form_version,
        )


# --- 3. Endorsement forms (DOCX) ----------------------------------------------------------------

ENDORSEMENT_TEXT = {
    "END-RA-01": (
        "Roadside Assistance Endorsement",
        "This endorsement adds 24-hour roadside assistance coverage, including towing, jump-starts, "
        "flat tire changes, and lockout service, up to $150 per occurrence, at no additional deductible. "
        "This endorsement does not modify any other coverage part of the Policy Wording.",
    ),
    "END-RS-01": (
        "Ride-share/TNC Endorsement",
        "This endorsement removes the Ride-share/TNC exclusion (Policy Wording Section 7.6, Standard "
        "Exclusions Addendum PAP-EXCL-2025 Exclusion 9) for the covered auto described on the "
        "Declarations page, while the named insured is logged into a transportation network company "
        "application in driver mode. All other terms, limits, and deductibles of the policy continue "
        "to apply unchanged during ride-share/TNC use.",
    ),
    "END-CE-01": (
        "Custom Equipment Endorsement",
        "This endorsement extends Comprehensive and Collision coverage (where selected) to custom "
        "equipment permanently installed in or on the covered auto, up to $2,500, subject to the "
        "otherwise-applicable deductible.",
    ),
    "END-NN-01": (
        "Named Non-Owner Household Driver Endorsement",
        "This endorsement extends Liability Coverage to a household resident driver who does not own "
        "a vehicle of their own, when operating the covered auto with permission, on the same terms "
        "as apply to the named insured.",
    ),
    "END-AI-01": (
        "Additional Insured Endorsement",
        "This endorsement adds the person or organization named on the Declarations page as an "
        "additional insured under Liability Coverage, solely with respect to their legal responsibility "
        "for the acts of the named insured while operating the covered auto with permission.",
    ),
    "END-GAP-01": (
        "Loan/Lease Gap Coverage Endorsement",
        "This endorsement pays the difference between a total loss settlement determined under the "
        "Total Loss Valuation Methodology and the amount owed on the covered auto's loan or lease, "
        "up to $10,000, when the settlement amount is less than the amount owed. This benefit is "
        "computed only after the base total-loss settlement figure is finalized and applies only to "
        "a total loss, not to a repairable claim.",
    ),
    "END-RR-01": (
        "Rental Reimbursement Increase Endorsement",
        "This endorsement increases the transportation expense limit shown on the Declarations page "
        "to $60 per day, maximum 45 days, replacing the standard limit described in Policy Wording "
        "Section 5.3 for the covered auto to which this endorsement is attached.",
    ),
    "END-OH-01": (
        "Ohio State Amendatory Endorsement",
        "This endorsement amends the Policy Wording to conform to Ohio statutory requirements where "
        "the base form's general provisions differ from state-specific requirements, including the "
        "cancellation and non-renewal notice periods described in the Underwriting Guidelines, "
        "Section 7. Where this endorsement conflicts with the base Policy Wording, this endorsement "
        "controls for policies garaged in Ohio.",
    ),
    "END-TEL-01": (
        "Telematics Program Enrollment Endorsement",
        "This endorsement enrolls the covered auto in Meridian Mutual's usage-based insurance "
        "program described in the Telematics and Usage-Based Insurance Program Terms. Enrollment "
        "affects premium calculation at renewal only; it does not modify coverage, limits, "
        "deductibles, or how a claim under this policy is adjudicated.",
    ),
    "END-UMPD-01": (
        "Uninsured Motorist Property Damage Endorsement",
        "This endorsement extends Uninsured Motorist Coverage (Policy Wording Part C) to property "
        "damage to the covered auto caused by an identified uninsured driver, up to $3,500, subject "
        "to a $250 deductible. This is separate from Collision Coverage and applies whether or not "
        "the covered auto carries Collision Coverage on the Declarations page.",
    ),
    "END-NVR-01": (
        "New Vehicle Replacement Cost Endorsement",
        "For a covered auto in its first model year of ownership that is declared a total loss "
        "under the Total Loss Valuation Methodology, this endorsement pays the cost of a new "
        "replacement vehicle of the same make and model, rather than the Actual Cash Value "
        "otherwise used for total loss settlement, subject to the endorsement's own $5,000 limit "
        "above the ACV figure.",
    ),
}


def build_endorsements() -> None:
    for code, (title, body) in ENDORSEMENT_TEXT.items():
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading(f"{facts.INSURER_NAME}", level=1)
        doc.add_heading(f"{title} (Form {code})", level=0)
        doc.add_paragraph(body)
        doc.add_paragraph("This endorsement is part of the policy and takes effect on the date shown on the Declarations page listing this endorsement.")

        slug = f"endorsement_{code.lower().replace('-', '_')}"
        docx_path = OUT / f"{slug}.docx"
        doc.save(str(docx_path))
        dst = CORPUS / "endorsements" / f"{slug}.docx"
        dst.write_bytes(docx_path.read_bytes())
        add_doc(slug, str(dst.relative_to(CORPUS)), "endorsements", "docx", title=f"{title} ({code})", formCode=code)


# --- 4. Repair estimates (DOCX, one per claim) --------------------------------------------------

def build_repair_estimates() -> None:
    for claim in facts.CLAIMS:
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading("Independent Repair Estimate", level=0)
        doc.add_paragraph(f"Claim Number: {claim.claim_number}")
        doc.add_paragraph(f"Policy Number: {claim.policy.policy_number}")
        doc.add_paragraph(f"Vehicle: {claim.policy.vehicle_year} {claim.policy.vehicle_make} {claim.policy.vehicle_model}")
        doc.add_paragraph(f"Date of Loss: {claim.date_of_loss}")
        doc.add_paragraph(f"Loss Type: {claim.loss_type}")

        doc.add_heading("Damage Description", level=2)
        doc.add_paragraph(claim.description)

        doc.add_heading("Estimate Summary", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Line Item", "Amount"
        cells = table.add_row().cells
        cells[0].text, cells[1].text = "Parts and Labor", f"${claim.estimated_damage:,}"
        cells = table.add_row().cells
        cells[0].text, cells[1].text = "Total Estimated Damage", f"${claim.estimated_damage:,}"

        slug = f"estimate_{claim.claim_number.lower().replace('-', '_')}"
        docx_path = OUT / f"{slug}.docx"
        doc.save(str(docx_path))
        dst = CORPUS / "claims" / f"{slug}.docx"
        dst.write_bytes(docx_path.read_bytes())
        add_doc(
            slug, str(dst.relative_to(CORPUS)), "claims", "docx",
            title=f"Repair Estimate — {claim.claim_number}",
            claimNumber=claim.claim_number, policyNumber=claim.policy.policy_number,
        )


# --- 5. Adjuster case notes (DOCX, one per claim) — narrates the five-step sequence -------------
#
# This mirrors the payout formula stated in the Claims Adjudication Guidelines
# (`payout = min(estimated_damage, applicable_limit) - applicable_deductible`, with the
# PAP-2025-STD glass-only waiver zeroing the deductible) purely to generate consistent corpus
# narrative — it is NOT the product's PayoutCalculationService, which is separate, unit-tested
# code built in Epic E.

def compute_narrative_payout(claim: facts.Claim) -> tuple[int | None, str]:
    ph = claim.policy
    if claim.flagged_anomaly:
        return None, "Not computed — claim flagged for review before adjudication (see below)."

    if claim.loss_type == "Liability":
        payout = min(claim.estimated_damage, ph.liability_pd)
        return payout, f"min(${claim.estimated_damage:,}, Property Damage limit ${ph.liability_pd:,}) = ${payout:,}. No deductible applies to third-party Liability claims."

    if claim.loss_type == "UM/UIM":
        payout = min(claim.estimated_damage, ph.um_uim_per_accident)
        return payout, f"min(${claim.estimated_damage:,}, UM/UIM per-accident limit ${ph.um_uim_per_accident:,}) = ${payout:,}. No deductible applies under Part C."

    if claim.loss_type == "Collision":
        deductible = ph.collision_deductible or 0
        payout = max(0, claim.estimated_damage - deductible)
        return payout, f"min(${claim.estimated_damage:,}, Collision coverage) - ${deductible:,} deductible = ${payout:,}."

    # Comprehensive
    waiver_applies = claim.is_glass_only and ph.form_version == facts.POLICY_FORM_V2 and claim.estimated_damage < 1_500
    deductible = 0 if waiver_applies else (ph.comprehensive_deductible or 0)
    payout = max(0, claim.estimated_damage - deductible)
    waiver_note = " Glass-only deductible waiver (PAP-2025-STD Section 5.4) applies." if waiver_applies else ""
    return payout, f"min(${claim.estimated_damage:,}, Comprehensive coverage) - ${deductible:,} deductible = ${payout:,}.{waiver_note}"


def build_case_notes() -> None:
    for claim in facts.CLAIMS:
        ph = claim.policy
        payout, computation_note = compute_narrative_payout(claim)

        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        doc.add_heading("Adjuster Case Notes", level=0)
        doc.add_paragraph(f"Claim Number: {claim.claim_number}")
        doc.add_paragraph(f"Policy Number: {ph.policy_number} ({ph.named_insured})")
        doc.add_paragraph(f"Governing Policy Form: {ph.form_version}")

        doc.add_heading("Step 1 — Policy Version Match", level=2)
        doc.add_paragraph(
            f"Date of loss {claim.date_of_loss} checked against the Declarations page for "
            f"{ph.policy_number}, effective {ph.effective_date} under Form {ph.form_version}. "
            + ("Date of loss falls within the policy period." if not claim.flagged_anomaly or "effective date" not in (claim.flagged_anomaly or "") else "ANOMALY: date of loss precedes the policy's effective date — see below.")
        )

        doc.add_heading("Step 2 — Coverage, Limits, and Deductibles", level=2)
        if claim.loss_type == "Collision":
            doc.add_paragraph(f"Collision coverage confirmed on Declarations page. Deductible: ${ph.collision_deductible or 0:,}." if ph.has_collision else "No Collision coverage on this policy — this claim cannot be adjudicated under Part D Collision.")
        elif claim.loss_type == "Comprehensive":
            doc.add_paragraph(f"Comprehensive coverage confirmed on Declarations page. Deductible: ${ph.comprehensive_deductible or 0:,}." if ph.has_comprehensive else "No Comprehensive coverage on this policy — this claim cannot be adjudicated under Part D Comprehensive.")
        elif claim.loss_type == "UM/UIM":
            doc.add_paragraph(f"Uninsured/Underinsured Motorist coverage confirmed on Declarations page. Per-accident limit: ${ph.um_uim_per_accident:,}. No deductible applies under Part C.")
        else:
            doc.add_paragraph(f"Liability coverage confirmed. Property Damage limit: ${ph.liability_pd:,}.")

        doc.add_heading("Step 3 — Exclusion Check", level=2)
        if claim.flagged_anomaly and "endorsement" in claim.flagged_anomaly.lower():
            doc.add_paragraph(claim.flagged_anomaly)
        elif "ride-share" in claim.description.lower() or "rideshare" in claim.description.lower():
            has_rs = any("Ride-share" in e for e in ph.endorsements)
            doc.add_paragraph(
                "Ride-share/TNC use confirmed in claim narrative. Ride-share/TNC Endorsement "
                + ("IS attached — exclusion does not apply." if has_rs else "is NOT attached — exclusion applies unless further information changes this.")
            )
        else:
            doc.add_paragraph("No exclusion indicators identified in the claim narrative against the Standard Exclusions Addendum.")

        doc.add_heading("Step 4 — Payout Computation", level=2)
        if payout is None:
            doc.add_paragraph(computation_note)
        else:
            doc.add_paragraph(f"Computed by the deterministic claims-calculation service: {computation_note}")

        doc.add_heading("Step 5 — Recommendation and Approval", level=2)
        if claim.flagged_anomaly:
            doc.add_paragraph(f"FLAGGED — routed to supervising adjuster before any recommendation is drafted. Reason: {claim.flagged_anomaly}")
        else:
            doc.add_paragraph(f"Recommendation drafted: approve payout of ${payout:,}, citing the coverage and exclusion sections referenced above. Pending adjuster approval before communication to policyholder or payment issuance.")

        slug = f"casenotes_{claim.claim_number.lower().replace('-', '_')}"
        docx_path = OUT / f"{slug}.docx"
        doc.save(str(docx_path))
        dst = CORPUS / "claims" / f"{slug}.docx"
        dst.write_bytes(docx_path.read_bytes())
        add_doc(
            slug, str(dst.relative_to(CORPUS)), "claims", "docx",
            title=f"Adjuster Case Notes — {claim.claim_number}",
            claimNumber=claim.claim_number, policyNumber=ph.policy_number,
        )


# --- 6. Scanned claim intake forms + incident reports (image-only PDF, T6 OCR target) -----------

def build_scanned_claim_forms() -> None:
    dummy_img = Image.new("L", (10, 10))
    draw = ImageDraw.Draw(dummy_img)

    for claim in facts.CLAIMS:
        ph = claim.policy
        lines = [
            f"{facts.INSURER_NAME}",
            "CLAIM INTAKE FORM",
            "",
            f"Claim Number: {claim.claim_number}",
            f"Policy Number: {ph.policy_number}",
            f"Named Insured: {ph.named_insured}",
            f"Address: {ph.address}",
            f"Vehicle: {ph.vehicle_year} {ph.vehicle_make} {ph.vehicle_model}, VIN {ph.vin}",
            f"Date of Loss: {claim.date_of_loss}",
            f"Time of Loss: reported as approximately mid-day unless otherwise noted",
            f"Location of Loss: see description below",
            f"Loss Type: {claim.loss_type}",
            "",
            "Description of Loss:",
            claim.description,
            "",
            "Vehicle Condition Prior to Loss:",
            "Insured reports no pre-existing damage to the areas affected by this loss.",
            "",
            "Witnesses:",
            "None identified at time of filing." if not claim.police_report_number else "See attached incident report summary for responding-officer account.",
            "",
            f"Police Report Number: {claim.police_report_number or 'N/A'}",
            "",
            "Other Insurance Involved: None disclosed" if claim.loss_type != "Liability" else "Third-party vehicle involved; see repair estimate for damage details.",
            "",
            "Insured's Statement:",
            f"\"{claim.description}\" — as relayed to the intake representative by {ph.named_insured}.",
            "",
            "Policyholder Signature: _______________________",
            f"Date Filed: {claim.date_of_loss}",
        ]
        pages = paginate(draw, lines)
        slug = f"intake_{claim.claim_number.lower().replace('-', '_')}"
        out_path = CORPUS / "claims" / f"{slug}.pdf"
        build_scanned_pdf(pages, out_path, seed=hash(claim.claim_number) % 1000)
        add_doc(
            slug, str(out_path.relative_to(CORPUS)), "claims", "pdf-scanned",
            title=f"Claim Intake Form (scanned) — {claim.claim_number}",
            claimNumber=claim.claim_number, policyNumber=claim.policy.policy_number,
            requiresOcr=True,
        )

    for claim in facts.CLAIMS:
        if not claim.police_report_number:
            continue
        lines = [
            "INCIDENT REPORT SUMMARY",
            f"Report Number: {claim.police_report_number}",
            f"Date: {claim.date_of_loss}",
            "",
            "Summary:",
            claim.description,
            "",
            "This is a synthetic incident report summary generated for a software assessment corpus.",
            "It does not correspond to any real law-enforcement record.",
        ]
        pages = paginate(draw, lines)
        slug = f"incident_{claim.police_report_number.lower().replace('-', '_')}"
        out_path = CORPUS / "claims" / f"{slug}.pdf"
        build_scanned_pdf(pages, out_path, seed=hash(claim.police_report_number) % 1000)
        add_doc(
            slug, str(out_path.relative_to(CORPUS)), "claims", "pdf-scanned",
            title=f"Incident Report Summary (scanned) — {claim.police_report_number}",
            claimNumber=claim.claim_number, requiresOcr=True,
        )


def main() -> None:
    reset_dirs()
    build_prose_docs()
    build_declarations()
    build_endorsements()
    build_repair_estimates()
    build_case_notes()
    build_scanned_claim_forms()

    manifest_path = CORPUS / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    print(f"Generated {len(manifest)} documents into {CORPUS}")
    by_format: dict[str, int] = {}
    for m in manifest:
        by_format[m["format"]] = by_format.get(m["format"], 0) + 1
    for fmt, count in sorted(by_format.items()):
        print(f"  {fmt}: {count}")


if __name__ == "__main__":
    sys.exit(main())
