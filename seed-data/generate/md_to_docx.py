"""Minimal Markdown -> DOCX renderer covering just what this corpus's content needs:
headings, paragraphs, bullet/numbered lists, simple pipe tables, and **bold** inline spans.
Not a general-purpose Markdown engine — deliberately small and easy to verify by reading it.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _add_inline(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def render(markdown_text: str, title: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    lines = markdown_text.strip("\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("|"):
            # Consume the whole table block (header, separator, rows).
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|[\s\-:|]+\|$", row)
            ]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        table.cell(r, c).text = cell
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)

        i += 1

    return doc


def write_docx(markdown_text: str, title: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    render(markdown_text, title).save(str(out_path))
